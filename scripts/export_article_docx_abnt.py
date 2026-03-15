from __future__ import annotations

import argparse
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DEFAULT_INPUT = Path("artigo/artigo_completo_export_docx.md")
DEFAULT_OUTPUT = Path("artigo/artigo_completo.docx")
AUTHOR_NAME = "Celso de Oliveira Lisboa"
DOC_TITLE = (
    "GKS: uma arquitetura de IA em dois estágios para problemas complexos "
    "de classificação: estudo de caso em detecção de intrusões"
)
KEYWORDS = (
    "inteligência artificial aplicada; classificação hierárquica; "
    "detecção de intrusões; ensemble learning; especialistas por classe"
)
ENGLISH_TITLE = (
    "GKS: A Two-Stage AI Architecture for Complex Classification Problems: "
    "A Case Study in Intrusion Detection"
)
FONT_FAMILY = "Times New Roman"


def _set_run_font(run, size_pt: float, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT_FAMILY
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FAMILY)


def _set_paragraph_format(
    paragraph,
    *,
    alignment: WD_ALIGN_PARAGRAPH | None,
    first_line_indent_cm: float | None,
    left_indent_cm: float | None = None,
    line_spacing: float = 1.5,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
) -> None:
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = None if first_line_indent_cm is None else Cm(first_line_indent_cm)
    fmt.left_indent = None if left_indent_cm is None else Cm(left_indent_cm)
    fmt.space_before = Pt(space_before_pt)
    fmt.space_after = Pt(space_after_pt)
    fmt.line_spacing = line_spacing


def _apply_runs(paragraph, size_pt: float, bold: bool | None = None, italic: bool | None = None) -> None:
    for run in paragraph.runs:
        _set_run_font(run, size_pt=size_pt, bold=bold, italic=italic)


def _contains_drawing(paragraph) -> bool:
    return any("w:drawing" in run._element.xml for run in paragraph.runs)


def _is_list_paragraph(paragraph) -> bool:
    ppr = paragraph._p.pPr
    return bool(ppr is not None and ppr.numPr is not None)


def _configure_styles(doc: Document) -> None:
    for style_name in ("Normal", "Title", "Author", "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Footnote Text"):
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = FONT_FAMILY
        style.font.size = Pt(12 if style_name != "Footnote Text" else 10)
        if style._element.rPr is not None:
            style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FAMILY)

    if "Normal" in doc.styles:
        fmt = doc.styles["Normal"].paragraph_format
        fmt.line_spacing = 1.5
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.first_line_indent = Cm(1.25)

    if "Footnote Text" in doc.styles:
        fmt = doc.styles["Footnote Text"].paragraph_format
        fmt.line_spacing = 1.0
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.first_line_indent = Cm(1.0)


def _remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _drop_duplicate_title_block(doc: Document) -> None:
    non_empty = [p for p in doc.paragraphs if p.text.strip()]
    if len(non_empty) < 4:
        return
    title_para, author_para, dup_title, dup_author = non_empty[:4]
    if title_para.style.name != "Title" or author_para.style.name != "Author":
        return
    if dup_title.text.strip() != title_para.text.strip():
        return
    if dup_author.text.strip() != author_para.text.strip():
        return
    _remove_paragraph(dup_title)
    _remove_paragraph(dup_author)


def _configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)


def _format_body_paragraph(paragraph) -> None:
    if _is_list_paragraph(paragraph):
        _set_paragraph_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent_cm=None,
            line_spacing=1.5,
        )
    else:
        _set_paragraph_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent_cm=1.25,
            line_spacing=1.5,
        )
    _apply_runs(paragraph, size_pt=12)


def _format_docx(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    doc.core_properties.title = DOC_TITLE
    doc.core_properties.author = AUTHOR_NAME
    doc.core_properties.keywords = KEYWORDS

    _configure_sections(doc)
    _configure_styles(doc)
    _drop_duplicate_title_block(doc)

    in_references = False
    awaiting_english_title = False

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if paragraph.style.name in {"Title", "Heading 1"}:
            _set_paragraph_format(
                paragraph,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                first_line_indent_cm=None,
                line_spacing=1.0,
                space_after_pt=12,
            )
            _apply_runs(paragraph, size_pt=12, bold=True)
            continue

        if paragraph.style.name == "Author" or text == AUTHOR_NAME:
            _set_paragraph_format(
                paragraph,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                first_line_indent_cm=None,
                line_spacing=1.0,
                space_after_pt=12,
            )
            _apply_runs(paragraph, size_pt=12)
            continue

        if text == "Referências":
            in_references = True
            awaiting_english_title = False

        if paragraph.style.name in {"Heading 2", "Heading 3", "Heading 4"}:
            _set_paragraph_format(
                paragraph,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent_cm=None,
                line_spacing=1.0,
                space_before_pt=12,
                space_after_pt=6,
            )
            _apply_runs(paragraph, size_pt=12, bold=True)
            awaiting_english_title = text == "Abstract"
            continue

        if awaiting_english_title and text:
            _set_paragraph_format(
                paragraph,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                first_line_indent_cm=None,
                line_spacing=1.0,
                space_after_pt=6,
            )
            _apply_runs(paragraph, size_pt=12, italic=True)
            awaiting_english_title = False
            continue

        if text.startswith("Palavras-chave:") or text.startswith("Keywords:"):
            _set_paragraph_format(
                paragraph,
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_line_indent_cm=None,
                line_spacing=1.0,
                space_after_pt=6,
            )
            _apply_runs(paragraph, size_pt=12)
            continue

        if in_references and text:
            _set_paragraph_format(
                paragraph,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent_cm=None,
                line_spacing=1.0,
                space_after_pt=12,
            )
            _apply_runs(paragraph, size_pt=12)
            continue

        if text.startswith("Figura "):
            _set_paragraph_format(
                paragraph,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                first_line_indent_cm=None,
                line_spacing=1.0,
                space_after_pt=6,
            )
            _apply_runs(paragraph, size_pt=10)
            continue

        if text.startswith("Tabela "):
            _set_paragraph_format(
                paragraph,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent_cm=None,
                line_spacing=1.0,
                space_after_pt=6,
            )
            _apply_runs(paragraph, size_pt=10, bold=True)
            continue

        if _contains_drawing(paragraph):
            _set_paragraph_format(
                paragraph,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                first_line_indent_cm=None,
                line_spacing=1.0,
                space_after_pt=6,
            )
            continue

        if not text:
            _set_paragraph_format(
                paragraph,
                alignment=None,
                first_line_indent_cm=None,
                line_spacing=1.0,
            )
            continue

        _format_body_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _set_paragraph_format(
                        paragraph,
                        alignment=paragraph.alignment,
                        first_line_indent_cm=None,
                        line_spacing=1.0,
                    )
                    _apply_runs(paragraph, size_pt=10)

    doc.save(str(docx_path))


def _run_pandoc(input_md: Path, output_docx: Path) -> None:
    cmd = [
        "pandoc",
        input_md.name,
        "-o",
        str(output_docx),
        "--from",
        "markdown+footnotes+pipe_tables+raw_attribute+fenced_divs",
        "--metadata",
        "lang=pt-BR",
        "--wrap=none",
    ]
    subprocess.run(cmd, cwd=str(input_md.parent), check=True)


def main() -> int:
    parser = argparse.ArgumentParser(description="Exporta artigo para DOCX com formatação ABNT-base.")
    parser.add_argument("--input", default=str(DEFAULT_INPUT), help="Markdown base de exportação.")
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT), help="DOCX de saída.")
    args = parser.parse_args()

    input_md = Path(args.input).resolve()
    output_docx = Path(args.output).resolve()

    if not input_md.exists():
        raise FileNotFoundError(f"Arquivo de entrada não encontrado: {input_md}")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    _run_pandoc(input_md=input_md, output_docx=output_docx)
    _format_docx(output_docx)
    print(f"[OK] DOCX exportado em {output_docx}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
